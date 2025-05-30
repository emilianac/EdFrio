# Função para converter números para texto por extenso
from num2words import num2words
from docx import Document
from flask import request
from babel.numbers import format_currency

def numero_por_extenso_com_decimal(n):
    # Aceita int ou float
    if not isinstance(n, (int, float)):
        return "Valor inválido"
    
    if n < 0 or n >= 1000000:
        return "Número fora do intervalo suportado (0-999999,99)"

    inteiro = int(n)
    
    # Cálculo do decimal - pega o valor decimal com duas casas, arredondando para evitar problemas
    decimal = int(round((n - inteiro) * 100))
    
    # Se decimal chegar a 100 (exemplo: 945.995 arredonda), ajusta para inteiro + 1 e decimal 0
    if decimal == 100:
        inteiro += 1
        decimal = 0

    texto_inteiro = numero_por_extenso(inteiro)
    
    if decimal > 0:
        texto_decimal = numero_por_extenso(decimal)
        return f"{texto_inteiro} vírgula {texto_decimal}"
    else:
        return texto_inteiro

# Função original ajustada para aceitar somente int
def numero_por_extenso(n):
    if not isinstance(n, int):
        return "Valor inválido"
    if n < 0 or n > 999999:
        return "Número fora do intervalo suportado (0-999999)"
    
    unidades = ["zero", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove"]
    especiais = ["dez", "onze", "doze", "treze", "quatorze", "quinze", "dezesseis", "dezessete", "dezoito", "dezenove"]
    dezenas = ["", "", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta", "oitenta", "noventa"]
    centenas = ["", "cento", "duzentos", "trezentos", "quatrocentos", "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"]

    def tres_digitos_por_extenso(num):
        if num == 0:
            return ""
        texto = ""
        c = num // 100
        d = (num % 100) // 10
        u = num % 10
        resto = num % 100

        if c > 0:
            if c == 1 and resto == 0:
                texto += "cem"
            else:
                texto += centenas[c]

        if resto > 0:
            if texto != "":
                texto += " e "
            if resto < 10:
                texto += unidades[resto]
            elif 10 <= resto < 20:
                texto += especiais[resto - 10]
            else:
                texto += dezenas[d]
                if u > 0:
                    texto += " e " + unidades[u]
        return texto

    milhares = n // 1000
    resto = n % 1000
    texto = ""

    if milhares > 0:
        if milhares == 1:
            texto += "mil"
        else:
            texto += tres_digitos_por_extenso(milhares) + " mil"

    if resto > 0:
        if texto != "":
            texto += " e "
        texto += tres_digitos_por_extenso(resto)
    elif n == 0:
        texto = "zero"

    return texto

# Função para converter números ordinais por extenso
def ordinal_por_extenso(n):
    ordinais_ate_20 = [
        "primeiro", "segundo", "terceiro", "quarto", "quinto",
        "sexto", "sétimo", "oitavo", "nono", "décimo",
        "décimo primeiro", "décimo segundo", "décimo terceiro",
        "décimo quarto", "décimo quinto", "décimo sexto",
        "décimo sétimo", "décimo oitavo", "décimo nono", "vigésimo"
    ]
    
    if not isinstance(n, int) or n < 1 or n > 31:
        return "Número fora do intervalo suportado"
    elif n <= 20:
        return ordinais_ate_20[n-1]
    elif 21 <= n <= 29:
        return "vigésimo " + ordinais_ate_20[n - 21]
    elif n == 30:
        return "trigésimo"
    elif n == 31:
        return "trigésimo primeiro"

# Dicionários para formatação de datas
meses = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril", 5: "maio", 6: "junho",
    7: "julho", 8: "agosto", 9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}

# Função de formatação de datas
def formatar_data(data_str):
    if data_str:
        try:
            if "-" in data_str:  # Formato yyyy-mm-dd
                ano, mes, dia = map(int, data_str.split('-'))
            elif "/" in data_str:  # Formato dd/mm/yyyy
                dia, mes, ano = map(int, data_str.split('/'))
            else:
                return data_str  # formato inesperado

            return f"{dia} de {meses[mes]} de {ano}"

        except (ValueError, KeyError):
            return data_str
    return ""

# Função para formatar o dia de pagamento
def formatar_dia_pagamento(dia_pag):
    if dia_pag is not None and dia_pag.isdigit():
        try:
            return num2words(int(dia_pag), to='ordinal', lang='pt')
        except NotImplementedError:
            return "Conversão ordinal não suportada para o idioma especificado"
    return dia_pag

# Função para substituir dados em um documento do Word
def substituir_dados_doc(doc, dados):
    def substituir_texto(elementos):
        for e in elementos:
            for run in e.runs:
                for chave, valor in dados.items():
                    if chave in run.text:
                        run.text = run.text.replace(chave, valor)

    # Substituir nos parágrafos
    substituir_texto(doc.paragraphs)

    # Substituir nas tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                substituir_texto(celula.paragraphs)

    return doc

VALOR_INVALIDO_MSG = "valor inválido"  # ou importe de outro lugar, se já tiver definido
def texto_entrada(valor_entrada_raw: str, tipo_entrada: str, numero_por_extenso) -> str:
    valor_entrada_raw = valor_entrada_raw.strip()
    
    if tipo_entrada != "sim":
        return "sem entrada"
    
    try:
        # Substitui vírgula por ponto para converter para float
        valor_float = float(valor_entrada_raw.replace(',', '.'))
    except ValueError:
        return VALOR_INVALIDO_MSG

    # Formatar valor para estilo brasileiro com 2 casas decimais
    valor_formatado = f"{valor_float:.2f}".replace('.', ',')

    # Para escrever por extenso, normalmente usamos a parte inteira
    valor_inteiro = int(valor_float)

    valor_extenso = numero_por_extenso(valor_inteiro)

    return f"com entrada de R$ {valor_formatado} ({valor_extenso} reais)"