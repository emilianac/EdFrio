from flask import Flask, render_template, request, send_file
from datetime import datetime
from io import BytesIO
from docx import Document
import os
from utils import formatar_data, formatar_dia_pagamento, numero_por_extenso, substituir_dados_doc, texto_entrada, numero_por_extenso_com_decimal, valor_monetario_para_float
from babel.numbers import format_currency

app = Flask(__name__)

VALOR_INVALIDO_MSG = "Valor inválido"
CAMINHO_MODELO = os.path.join(os.path.dirname(__file__), 'CONTRATO DE LOCAÇÃO DE RESFRIADOR.docx')
CAMINHO_MODELO_ORDENHA = os.path.join(os.path.dirname(__file__), 'CONTRATO DE LOCAÇÃO DE ORDENHA.docx')
CAMINHO_MODELO_VENDA = os.path.join(os.path.dirname(__file__), 'CONTRATO DE COMPRA E VENDA COM RESERVA DE DOMÍNIO.docx')

# Página principal
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/form_aluguel', methods=['GET', 'POST'])
def form_aluguel():
    if request.method == 'POST':

        # Formatação CPF e RG
        cpf = request.form.get("cpf", "").zfill(11)
        cpf_formatado = f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}" if len(cpf) == 11 else cpf
        rg = request.form.get("num_rg", "").zfill(9)
        rg_formatado = f"{rg[:2]}.{rg[2:5]}.{rg[5:8]}-{rg[8]}" if len(rg) == 9 else rg

        # Formatação das datas
        data_inicio = request.form.get("data_inicio", "").strip()
        data_fim = request.form.get("data_fim", "").strip()

        data_inicio_formatada = formatar_data(data_inicio)
        if data_inicio_formatada == data_inicio or data_inicio_formatada == "":  # If the format function returns the original string, it indicates an invalid date format
            data_inicio_formatada = "Data inválida"  # Fallback value for invalid dates
        data_fim_formatada = formatar_data(data_fim)
        if data_fim_formatada == data_fim or data_fim_formatada == "":  # Check for invalid formats
            data_fim_formatada = "Data inválida"  # Fallback value for invalid dates
        if data_fim_formatada == data_fim or not data_fim_formatada:  # If the format function returns the original string or empty
            data_fim_formatada = "Data inválida"  # Fallback value for invalid dates
        
        # Cache the formatted current date
        data_atual_formatada = formatar_data(datetime.now().strftime("%d/%m/%Y"))

        # Função para obter os dados do formulário
        def obter_dados_formulario():
            prazo = request.form.get("prazo_meses", "")
            dia_pag = request.form.get("dia_pag", "")
            valor = request.form.get("valor", "")
            quantidade = request.form.get("quantidade", "")
            multa = request.form.get("multa", "")
            tipo_rg = request.form.get("rg_tipo", "")
            transformador = request.form.get("transformador", "")
            qtd_trans = request.form.get("qtd_trans", "")
            qtd_va = request.form.get("qtd_va", "")

            if transformador == "trans_sim":
                if qtd_trans.isdigit() and 1 <= int(qtd_trans) <= 31:
                    qtd_trans_ext = numero_por_extenso(int(qtd_trans))
                else:
                    qtd_trans_ext = VALOR_INVALIDO_MSG

                info_trans = f" com {qtd_trans} ({qtd_trans_ext}) TRANSFORMADOR {qtd_va} VA"
            else:
                info_trans = ""

            if tipo_rg == "rg_antigo":
                rg_valor = ", RG " + rg_formatado
                orgao_emissor_valor = request.form.get("orgao_emissor", "")
                estado_emissor_valor = request.form.get("estado_emissor", "")
            else:
                rg_valor = ""
                orgao_emissor_valor = ""
                estado_emissor_valor = ""
            return {
                "<<NOME>>": request.form.get("nome_completo", ""),
                "<<CPF>>": cpf_formatado,
                "<<RG>>": rg_valor,
                "<<ORGAO EMISSOR>>": orgao_emissor_valor,
                "<<ESTADO EMISSOR>>": estado_emissor_valor,
                "<<CIDADE>>": request.form.get("cidade", ""),
                "<<ESTADO>>": request.form.get("estado", ""),
                "<<LOGRADOURO>>": request.form.get("logradouro", ""),
                "<<NUM>>": request.form.get("numero", ""),
                "<<BAIRRO>>": request.form.get("bairro", ""),
                "<<CEP>>": request.form.get("cep", ""),
                "<<QUANTIDADE>>": request.form.get("quantidade", ""),
                "<<QUANTIDADE_EXTENSO>>": numero_por_extenso(int(quantidade)) if quantidade.isdigit() and 1 <= int(quantidade) <= 31 else VALOR_INVALIDO_MSG,
                "<<MARCA>>": request.form.get("marca", ""),
                "<<MODELO>>": request.form.get("modelo", ""),
                "<<TRANSFORMADOR>>": info_trans,
                "<<PRAZO>>": prazo,
                "<<PRAZO_EXTENSO>>": numero_por_extenso(int(prazo)) if prazo.isdigit() and 1 <= int(prazo) <= 31 else VALOR_INVALIDO_MSG,
                "<<DIA PAG>>": request.form.get("dia_pag", ""),
                "<<DIA_ORD>>": formatar_dia_pagamento(dia_pag) if dia_pag.isdigit() and 1 <= int(dia_pag) <= 31 else VALOR_INVALIDO_MSG,
                "<<BEGIN>>": data_inicio_formatada,
                "<<FIM>>": data_fim_formatada,
                "<<VALOR>>": valor,
                "<<VALOR_EXTENSO>>": numero_por_extenso(int(valor)),
                "<<MULTA>>": multa,
                "<<MULTA_EXTENSO>>": numero_por_extenso(int(multa)) if multa.isdigit() else VALOR_INVALIDO_MSG,
                "<<DATA>>": data_atual_formatada,
            }
        
        # Verifica se o modelo de contrato existe
        if not os.path.isfile(CAMINHO_MODELO):
            return "Modelo de contrato não encontrado.", 500

        # Carrega o modelo do documento
        doc = Document(CAMINHO_MODELO)

        # Obtenha os dados do formulário
        dados = obter_dados_formulario()

        # Função de substituição
        def substituir_texto(elementos):
            replacements = {k: v for k, v in dados.items() if k in ''.join([run.text for e in elementos for run in e.runs])}
            for e in elementos:
                for run in e.runs:
                    for chave, valor in replacements.items():
                        run.text = run.text.replace(chave, valor)

        # Substitui nos parágrafos
        substituir_texto(doc.paragraphs)

        # Substitui nos parágrafos das tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    substituir_texto(celula.paragraphs)
        substituir_texto(doc.paragraphs)

        # Substitui nas tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    substituir_texto(celula.paragraphs)
                    
        # Define o nome do arquivo para download
        nome_pessoa = request.form.get("nome_completo", "documento").strip().replace(" ", "_")
        nome_arquivo = f"CONTRATO DE LOCAÇÃO RESFRIADOR_{nome_pessoa}.docx"

        # Salva no stream
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Retorna para download no Flask
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=nome_arquivo,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    # Renderiza o template do formulário
    return render_template('form_aluguel.html')

@app.route('/form_aluguel_ord', methods=['GET', 'POST'])
def form_aluguel_ord():
    if request.method == 'POST':

        # Formatação CPF e RG
        cpf = request.form.get("cpf", "").zfill(11)
        cpf_formatado = f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}" if len(cpf) == 11 else cpf
        rg = request.form.get("num_rg", "").zfill(9)
        rg_formatado = f"{rg[:2]}.{rg[2:5]}.{rg[5:8]}-{rg[8]}" if len(rg) == 9 else rg

        # Formatação das datas
        data_inicio = request.form.get("data_inicio", "").strip()
        data_fim = request.form.get("data_fim", "").strip()

        data_inicio_formatada = formatar_data(data_inicio)
        if data_inicio_formatada == data_inicio or data_inicio_formatada == "":  # If the format function returns the original string, it indicates an invalid date format
            data_inicio_formatada = "Data inválida"  # Fallback value for invalid dates
        data_fim_formatada = formatar_data(data_fim)
        if data_fim_formatada == data_fim or data_fim_formatada == "":  # Check for invalid formats
            data_fim_formatada = "Data inválida"  # Fallback value for invalid dates
        if data_fim_formatada == data_fim or not data_fim_formatada:  # If the format function returns the original string or empty
            data_fim_formatada = "Data inválida"  # Fallback value for invalid dates
        
        # Cache the formatted current date
        data_atual_formatada = formatar_data(datetime.now().strftime("%d/%m/%Y"))

        # Função para obter os dados do formulário
        def obter_dados_formulario():
            prazo = request.form.get("prazo_meses", "")
            dia_pag = request.form.get("dia_pag", "")
            valor = request.form.get("valor", "")
            quantidade = request.form.get("quantidade", "")
            qtde = request.form.get("qtd", "")
            multa = request.form.get("multa", "")
            tipo_rg = request.form.get("rg_tipo", "")
            
            if tipo_rg == "rg_antigo":
                rg_valor = ", RG " + rg_formatado
                orgao_emissor_valor = request.form.get("orgao_emissor", "")
                estado_emissor_valor = request.form.get("estado_emissor", "")
            else:
                rg_valor = ""
                orgao_emissor_valor = ""
                estado_emissor_valor = ""

            return {
                "<<NOME>>": request.form.get("nome_completo", ""),
                "<<CPF>>": cpf_formatado,
                "<<RG>>": rg_valor,
                "<<ORGAO EMISSOR>>": orgao_emissor_valor,
                "<<ESTADO EMISSOR>>": estado_emissor_valor,
                "<<CIDADE>>": request.form.get("cidade", ""),
                "<<ESTADO>>": request.form.get("estado", ""),
                "<<LOGRADOURO>>": request.form.get("logradouro", ""),
                "<<BAIRRO>>": request.form.get("bairro", ""),
                "<<CEP>>": request.form.get("cep", ""),
                "<<NUM>>": request.form.get("numero", ""),
                "<<QUANTIDADE>>": request.form.get("quantidade", ""),
                "<<QUANTIDADE_EXTENSO>>": numero_por_extenso(int(quantidade)) if quantidade.isdigit() and 1 <= int(quantidade) <= 31 else VALOR_INVALIDO_MSG,
                "<<OBJETO>>": request.form.get("objeto", ""),
                "<<MARCA>>": request.form.get("marca", ""),
                "<<MODELO>>": request.form.get("modelo", ""),
                "<<CAVALO>>": request.form.get("cavalo", ""),
                "<<QTD>>": request.form.get("qtd", ""),
                "<<QTD_EXTENSO>>": numero_por_extenso(int(qtde)) if qtde.isdigit() and 1 <= int(qtde) <= 31 else VALOR_INVALIDO_MSG,
                "<<PRAZO>>": prazo,
                "<<PRAZO_EXTENSO>>": numero_por_extenso(int(prazo)) if prazo.isdigit() and 1 <= int(prazo) <= 31 else VALOR_INVALIDO_MSG,
                "<<DIA PAG>>": request.form.get("dia_pag", ""),
                "<<DIA_ORD>>": formatar_dia_pagamento(dia_pag) if dia_pag.isdigit() and 1 <= int(dia_pag) <= 31 else VALOR_INVALIDO_MSG,
                "<<INICIO>>": data_inicio_formatada,
                "<<FIM>>": data_fim_formatada,
                "<<VALOR>>": valor,
                "<<VALOR_EXTENSO>>": numero_por_extenso(int(valor)),
                "<<MULTA>>": multa,
                "<<MULTA_EXTENSO>>": numero_por_extenso(int(multa)) if multa.isdigit() else VALOR_INVALIDO_MSG,
                "<<DATA>>": data_atual_formatada,
            }
        
        # Verifica se o modelo de contrato existe
        if not os.path.isfile(CAMINHO_MODELO_ORDENHA):
            return "Modelo de contrato não encontrado.", 500

        # Carrega o modelo do documento
        doc = Document(CAMINHO_MODELO_ORDENHA)

        # Obtenha os dados do formulário
        dados = obter_dados_formulario()

        substituir_dados_doc(doc, dados)
        # Define o nome do arquivo para download
        nome_pessoa = request.form.get("nome_completo", "documento").strip().replace(" ", "_")
        nome_arquivo = f"CONTRATO DE LOCAÇÃO ORDENHA_{nome_pessoa}.docx"

        # Salva no stream
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Retorna para download no Flask
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=nome_arquivo,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    # Renderiza o template do formulário
    return render_template('form_aluguel_ord.html')


@app.route('/form_venda', methods=['GET', 'POST'])
def form_venda():
    if request.method == 'POST':

        # Formatação CPF e RG
        cpf = request.form.get("cpf", "").zfill(11)
        cpf_formatado = f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}" if len(cpf) == 11 else cpf

        # Formatação das datas
        data_inicio = request.form.get("data_inicio", "").strip()
        data_fim = request.form.get("data_fim", "").strip()

        data_inicio_formatada = formatar_data(data_inicio)
        if data_inicio_formatada == data_inicio or data_inicio_formatada == "":  # If the format function returns the original string, it indicates an invalid date format
            data_inicio_formatada = "Data inválida"  # Fallback value for invalid dates
        data_fim_formatada = formatar_data(data_fim)
        if data_fim_formatada == data_fim or data_fim_formatada == "":  # Check for invalid formats
            data_fim_formatada = "Data inválida"  # Fallback value for invalid dates
        if data_fim_formatada == data_fim or not data_fim_formatada:  # If the format function returns the original string or empty
            data_fim_formatada = "Data inválida"  # Fallback value for invalid dates
        
        # Cache the formatted current date
        data_atual_formatada = formatar_data(datetime.now().strftime("%d/%m/%Y"))

        # Função para obter os dados do formulário
        def obter_dados_formulario():
            valor = request.form.get("valor", "")
            valor_parcela_float = valor_monetario_para_float(valor)
            quantidade_res = request.form.get("quantidade_res", "")
            quantidade_ord = request.form.get("quantidade_ord", "")
            tipo_venda = request.form.get("tipo", "")
            prazo = request.form.get("prazo_meses", "")
            valor_total = request.form.get("valor_total", "")
            tipo_entrada = request.form.get("tipo_entrada", "")
            valor_entrada = request.form.get("valor_entrada", "")
            try:
                # Converter para float, tratando pontos como separador decimal
                valor_float = float(valor_entrada.replace(',', '.'))  
                # Formatar para string no formato '00,00'
                valor_formatado = f"{valor_float:,.2f}".replace('.', ',')
            except ValueError:
                valor_formatado = "0,00"  # valor padrão caso a conversão falhe

            tipo_entrada = request.form.get("tipo_entrada", "")
            texto = texto_entrada(valor_formatado, tipo_entrada, numero_por_extenso_com_decimal)

            if tipo_venda == "venda_res":
                return {
                    "<<NOME>>": request.form.get("nome_completo", ""),
                    "<<CPF>>": cpf_formatado,
                    "<<CIDADE>>": request.form.get("cidade", ""),
                    "<<ESTADO>>": request.form.get("estado", ""),
                    "<<LOGRADOURO>>": request.form.get("logradouro", ""),
                    "<<BAIRRO>>": request.form.get("bairro", ""),
                    "<<CEP>>": request.form.get("cep", ""),
                    "<<NUMERO>>": request.form.get("numero", ""),
                    "<<QUANTIDADE>>": request.form.get("quantidade_res", ""),
                    "<<QUANTIDADE_EXTENSO>>": numero_por_extenso(int(quantidade_res)) if quantidade_res.isdigit() and 1 <= int(quantidade_res) <= 31 else VALOR_INVALIDO_MSG,
                    "<<OBJETO>>": "Tanque Resfriador",
                    "<<MARCA>>": request.form.get("marca_res", ""),
                    "<<MODELO>>": request.form.get("modelo_res", ""),
                    "<<COMPLEMENTO>>": "",
                    "<<VALOR>>": valor_total,
                    "<<VALOR_EXTENSO>>": numero_por_extenso(int(valor_total)),
                    "<<ENTRADA>>": texto,
                    "<<PARCELA>>": prazo,
                    "<<PARCELA_EXTENSO>>": numero_por_extenso(int(prazo)) if prazo.isdigit() and 1 <= int(prazo) <= 31 else VALOR_INVALIDO_MSG,
                    "<<VALOR_PARCELA>>": valor,
                    "<<VALOR_PARCELA_EXTENSO>>": numero_por_extenso_com_decimal(valor_parcela_float) if valor_parcela_float is not None else VALOR_INVALIDO_MSG,
                    "<<INICIO>>": data_inicio_formatada,
                    "<<FIM>>": data_fim_formatada,
                    "<<DATA>>": data_atual_formatada,
                }
            else:
                return {
                    "<<NOME>>": request.form.get("nome_completo", ""),
                    "<<CPF>>": cpf_formatado,
                    "<<CIDADE>>": request.form.get("cidade", ""),
                    "<<ESTADO>>": request.form.get("estado", ""),
                    "<<LOGRADOURO>>": request.form.get("logradouro", ""),
                    "<<BAIRRO>>": request.form.get("bairro", ""),
                    "<<CEP>>": request.form.get("cep", ""),
                    "<<NUMERO>>": request.form.get("numero", ""),
                    "<<QUANTIDADE>>": request.form.get("quantidade_ord", ""),
                    "<<QUANTIDADE_EXTENSO>>": numero_por_extenso(int(quantidade_ord)) if quantidade_ord.isdigit() and 1 <= int(quantidade_ord) <= 31 else VALOR_INVALIDO_MSG,
                    "<<OBJETO>>": "Ordenhadeira",
                    "<<MARCA>>": request.form.get("marca_ord", ""),
                    "<<MODELO>>": request.form.get("modelo_ord", ""),
                    "<<COMPLEMENTO>>": request.form.get("complemento_ord", ""),
                    "<<VALOR>>": valor_total,
                    "<<VALOR_EXTENSO>>": numero_por_extenso(int(valor_total)),
                    "<<ENTRADA>>": texto,
                    "<<PARCELA>>": prazo,
                    "<<PARCELA_EXTENSO>>": numero_por_extenso(int(prazo)) if prazo.isdigit() and 1 <= int(prazo) <= 31 else VALOR_INVALIDO_MSG,
                    "<<VALOR_PARCELA>>": valor,
                    "<<VALOR_PARCELA_EXTENSO>>": numero_por_extenso_com_decimal(valor_parcela_float) if valor_parcela_float is not None else VALOR_INVALIDO_MSG,
                    "<<INICIO>>": data_inicio_formatada,
                    "<<FIM>>": data_fim_formatada,
                    "<<DATA>>": data_atual_formatada,
                }
            
           
        # Verifica se o modelo de contrato existe
        if not os.path.isfile(CAMINHO_MODELO_VENDA):
            return "Modelo de contrato não encontrado.", 500

        # Carrega o modelo do documento
        doc = Document(CAMINHO_MODELO_VENDA)

        # Obtenha os dados do formulário
        dados = obter_dados_formulario()

        substituir_dados_doc(doc, dados)
        # Define o nome do arquivo para download
        nome_pessoa = request.form.get("nome_completo", "documento").strip().replace(" ", "_")
        nome_arquivo = f"CONTRATO DE VENDA_{nome_pessoa}.docx"

        # Salva no stream
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Retorna para download no Flask
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=nome_arquivo,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    # Renderiza o template do formulário
    return render_template('form_venda.html')

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
# This code is a Flask application that serves a form for generating a rental contract.
# It processes the form data, formats it, and generates a Word document based on a template.